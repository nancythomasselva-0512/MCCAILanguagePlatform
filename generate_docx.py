import markdown
from docx import Document
from docx.shared import Pt
import re

def create_docx():
    with open('MCC_AI_Language_Platform_Enterprise_Documentation.md', 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = Document()
    doc.add_heading('MCC AI Language Platform - Enterprise Documentation', 0)
    
    # Simple manual parse of markdown
    lines = md_text.split('\n')
    
    in_table = False
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('|') and '---' in line:
            pass # table separator
        elif line.startswith('|'):
            # simple table row, just add as text for now
            doc.add_paragraph(line.replace('|', '  ').strip())
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            # Handle basic bold parsing
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    doc.save('MCC_AI_Language_Platform_Enterprise_Documentation.docx')
    print("Docx created successfully.")

if __name__ == '__main__':
    create_docx()
